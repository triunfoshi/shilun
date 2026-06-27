"""波浪/技术分析外部文档解析器。

从用户上传的 .docx/.pdf 里提取大盘关键价位（支撑、压力、波浪标记），
作为人工标注位叠加到 PART1 的自动计算结果上。

策略：
1. 先用正则提取明显的"支撑/压力/目标位 + 数字"组合
2. 同时记录原文上下文，便于人工核对
3. AI 解读（Claude/OpenAI）作为可选后端，未配置时降级为纯正则
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any


# ── 关键词配置 ──────────────────────────────────────────────────────────────
# 注意：A股大盘点位通常在 1000~6000 之间，过滤掉明显不是点位的数字（如年份、百分比）
_PRICE_RE = r"(\d{3,4}(?:\.\d{1,2})?)"  # 3-4 位整数，可选 1-2 位小数

_LEVEL_KEYWORDS = {
    "support": [
        "支撑位", "支撑", "下方支撑", "强支撑", "关键支撑", "重要支撑",
        "防守位", "防守点", "止跌位",
    ],
    "pressure": [
        "压力位", "压力", "上方压力", "强压力", "关键压力", "重要压力",
        "阻力位", "阻力", "目标位", "目标点", "上方目标",
    ],
    "wave": [
        "波", "浪", "Wave", "wave", "推动浪", "调整浪",
        "A浪", "B浪", "C浪", "1浪", "2浪", "3浪", "4浪", "5浪",
    ],
    "breakdown": ["破位", "跌破", "失守"],
    "breakout": ["突破", "站上", "收复"],
}


@dataclass
class ExtractedLevel:
    """一个提取出的关键价位。"""

    price: float
    kind: str  # "support" | "pressure" | "wave" | "neutral"
    label: str  # 描述（如"上方压力"）
    context: str  # 原文上下文（前后 30 字）
    confidence: float = 0.5  # 0-1
    source_doc: str = ""

    def to_dict(self) -> dict[str, Any]:
        return {
            "price": self.price,
            "kind": self.kind,
            "label": self.label,
            "context": self.context,
            "confidence": self.confidence,
            "source_doc": self.source_doc,
        }


@dataclass
class WaveDocResult:
    """文档解析结果。"""

    source_doc: str
    text_length: int
    levels: list[ExtractedLevel] = field(default_factory=list)
    ai_summary: str = ""  # AI 解读的整体结论（可选）
    ai_provider: str = ""  # "anthropic" | "openai" | "" (纯正则)
    benchmark_ticker: str | None = None  # 文档主要讨论的指数
    raw_text_preview: str = ""

    def to_dict(self) -> dict[str, Any]:
        return {
            "source_doc": self.source_doc,
            "text_length": self.text_length,
            "levels": [lv.to_dict() for lv in self.levels],
            "ai_summary": self.ai_summary,
            "ai_provider": self.ai_provider,
            "benchmark_ticker": self.benchmark_ticker,
            "raw_text_preview": self.raw_text_preview,
            "level_count": len(self.levels),
        }


# ── 文档读取 ────────────────────────────────────────────────────────────────
def extract_text_from_docx(file_bytes: bytes) -> str:
    """从 .docx 文件读取所有文本（段落 + 表格）。"""
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """从 .pdf 文件读取所有文本（pypdf）。"""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    parts: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            text = ""
        if text.strip():
            parts.append(text.strip())
    return "\n".join(parts)


def extract_text(filename: str, file_bytes: bytes) -> str:
    """根据文件扩展名分发到对应解析器。"""
    name_lower = filename.lower()
    if name_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    if name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if name_lower.endswith((".txt", ".md")):
        return file_bytes.decode("utf-8", errors="ignore")
    raise ValueError(f"不支持的文件格式：{filename}。当前支持 .docx / .pdf / .txt / .md。")


# ── 正则提取 ────────────────────────────────────────────────────────────────
def _classify_keyword(window: str) -> tuple[str, str]:
    """判断窗口文本里出现的关键词类别。返回 (kind, label)。"""
    for kind, keywords in _LEVEL_KEYWORDS.items():
        for kw in keywords:
            if kw in window:
                return kind, kw
    return "neutral", ""


def extract_levels_by_regex(text: str, source_doc: str = "") -> list[ExtractedLevel]:
    """用正则在文本中扫描"关键词 + 数字"组合，提取关键价位。

    扫描方式：
    - 找所有 3-4 位数字（A 股大盘合理点位范围）
    - 取数字前后 ±30 字符窗口
    - 看窗口里有没有 _LEVEL_KEYWORDS 关键词
    """
    levels: list[ExtractedLevel] = []
    seen_prices: set[float] = set()  # 同一价位只保留第一次出现
    for match in re.finditer(_PRICE_RE, text):
        try:
            price = float(match.group(1))
        except ValueError:
            continue
        # 大盘点位过滤：A股主流指数在 800~6500 之间
        if price < 800 or price > 6500:
            continue
        start = max(0, match.start() - 30)
        end = min(len(text), match.end() + 30)
        context = text[start:end].replace("\n", " ")
        kind, label = _classify_keyword(context)
        if kind == "neutral":
            continue  # 没有关键词的数字忽略
        if price in seen_prices:
            continue
        seen_prices.add(price)
        # 置信度：关键词靠近数字 + 上下文里有波浪/趋势相关词，则提高
        confidence = 0.5
        distance = abs(context.find(label) - context.find(match.group(1)))
        if distance <= 10:
            confidence += 0.2
        if any(w in context for w in ("浪", "波", "Wave", "wave")):
            confidence += 0.15
        confidence = min(1.0, confidence)
        levels.append(
            ExtractedLevel(
                price=price,
                kind=kind,
                label=label,
                context=context.strip(),
                confidence=round(confidence, 2),
                source_doc=source_doc,
            )
        )
    # 按 kind + price 排序
    return sorted(levels, key=lambda lv: (lv.kind, lv.price))


# ── AI 解读（可选）──────────────────────────────────────────────────────────
def call_ai_summary(text: str, provider: str = "anthropic") -> tuple[str, str]:
    """让 AI 给整篇文档输出"核心结论"。

    返回 (summary_text, provider_used)。
    任何失败均降级为 ("", "")，由调用方决定是否提示用户。
    """
    if not text.strip():
        return "", ""
    text_clipped = text[:6000]  # 控制 token，黑兔单日文档通常 3-5k 字
    prompt = (
        "你是 A 股技术分析助手。下面是一篇大盘分析文档的全文。请用 3-5 句话总结：\n"
        "1) 文档对明日大盘的整体判断（多/空/震荡）；\n"
        "2) 列出文档明确给出的支撑位、压力位、波浪关键位（如有）；\n"
        "3) 文档提到的核心风险或机会。\n"
        "不要复述原文，给出结构化结论。\n\n"
        f"文档内容：\n{text_clipped}"
    )
    if provider == "anthropic":
        try:
            import os
            import anthropic  # type: ignore
            if not os.environ.get("ANTHROPIC_API_KEY"):
                return "", ""
            client = anthropic.Anthropic()
            msg = client.messages.create(
                model="claude-haiku-4-5",
                max_tokens=600,
                messages=[{"role": "user", "content": prompt}],
            )
            return msg.content[0].text.strip(), "anthropic"
        except Exception:  # noqa: BLE001
            return "", ""
    if provider == "openai":
        try:
            import os
            from openai import OpenAI  # type: ignore
            if not os.environ.get("OPENAI_API_KEY"):
                return "", ""
            client = OpenAI()
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=600,
            )
            return resp.choices[0].message.content.strip(), "openai"
        except Exception:  # noqa: BLE001
            return "", ""
    return "", ""


# ── 主入口 ──────────────────────────────────────────────────────────────────
def parse_wave_document(
    filename: str,
    file_bytes: bytes,
    *,
    use_ai: bool = False,
    ai_provider: str = "anthropic",
) -> WaveDocResult:
    """主入口：解析一份文档，返回结构化结果。"""
    text = extract_text(filename, file_bytes)
    if not text.strip():
        return WaveDocResult(
            source_doc=filename,
            text_length=0,
            raw_text_preview="（文档未提取到文本，可能是扫描件或图片型 PDF）",
        )

    levels = extract_levels_by_regex(text, source_doc=filename)
    ai_summary = ""
    ai_provider_used = ""
    if use_ai:
        ai_summary, ai_provider_used = call_ai_summary(text, provider=ai_provider)

    preview = text[:300].replace("\n", " ") + ("..." if len(text) > 300 else "")
    return WaveDocResult(
        source_doc=filename,
        text_length=len(text),
        levels=levels,
        ai_summary=ai_summary,
        ai_provider=ai_provider_used,
        raw_text_preview=preview,
    )
